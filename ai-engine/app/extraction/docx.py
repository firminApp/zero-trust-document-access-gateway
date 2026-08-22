"""Extraction DOCX : paragraphes, tableaux, en-têtes et pieds de page.

Un formulaire d'enrôlement met souvent l'identité dans un tableau et le nom du
client dans l'en-tête : ne lire que `document.paragraphs` laisserait ces
données invisibles à la détection, donc en clair à la restitution.
"""

from __future__ import annotations

import io
import logging

from docx import Document

from app.extraction.resultat import ResultatExtraction, depuis_texte
from app.models import MethodeExtraction

logger = logging.getLogger(__name__)


def extraire(contenu: bytes) -> ResultatExtraction:
    document = Document(io.BytesIO(contenu))
    lignes: list[str] = []

    for section in document.sections:
        lignes.extend(_lignes_conteneur(section.header))
        lignes.extend(_lignes_conteneur(section.footer))

    for paragraphe in document.paragraphs:
        if paragraphe.text.strip():
            lignes.append(paragraphe.text)

    for tableau in document.tables:
        for rangee in tableau.rows:
            cellules = [cellule.text.strip() for cellule in rangee.cells]
            if any(cellules):
                lignes.append(" | ".join(cellules))

    texte = "\n".join(lignes)
    return depuis_texte(texte, MethodeExtraction.docx, pages_brutes=[texte])


def _lignes_conteneur(conteneur) -> list[str]:  # noqa: ANN001 - type docx interne
    lignes: list[str] = []
    try:
        for paragraphe in conteneur.paragraphs:
            if paragraphe.text.strip():
                lignes.append(paragraphe.text)
        for tableau in conteneur.tables:
            for rangee in tableau.rows:
                cellules = [cellule.text.strip() for cellule in rangee.cells]
                if any(cellules):
                    lignes.append(" | ".join(cellules))
    except Exception as exc:  # pragma: no cover - en-tête absent
        logger.debug("En-tête/pied illisible : %s", exc)
    return lignes
