"""Fixtures : un document de chaque format pris en charge."""

from __future__ import annotations

import io

import pytest

TEXTE_REFERENCE = (
    "ATTESTATION DE DOMICILE\n"
    "Je soussigne Mamadou FALL, ne le 03/07/1988 a Thies,\n"
    "titulaire de la CNI n 1988070312345,\n"
    "courriel mamadou.fall@exemple.sn, telephone +221 77 555 44 33,\n"
    "IBAN SN91SN0100152000048500000765.\n"
)


@pytest.fixture(scope="session")
def texte_reference() -> str:
    return TEXTE_REFERENCE


@pytest.fixture(scope="session")
def pdf_natif() -> bytes:
    """PDF avec couche texte — extraction directe attendue."""
    import fitz

    document = fitz.open()
    page = document.new_page()
    page.insert_text((60, 80), TEXTE_REFERENCE, fontsize=11, fontname="helv")
    octets = document.tobytes()
    document.close()
    return octets


@pytest.fixture(scope="session")
def pdf_scanne(image_jpeg: bytes) -> bytes:
    """PDF sans couche texte — doit basculer vers l'OCR."""
    import fitz

    document = fitz.open()
    page = document.new_page(width=595, height=842)
    page.insert_image(fitz.Rect(0, 0, 595, 842), stream=image_jpeg)
    octets = document.tobytes()
    document.close()
    return octets


@pytest.fixture(scope="session")
def docx_simple() -> bytes:
    """DOCX avec en-tête, paragraphes et tableau."""
    from docx import Document

    document = Document()
    document.sections[0].header.paragraphs[0].text = "Dossier client GZ-889201"
    for ligne in TEXTE_REFERENCE.strip().split("\n"):
        document.add_paragraph(ligne)

    tableau = document.add_table(rows=2, cols=2)
    tableau.cell(0, 0).text = "Courriel"
    tableau.cell(0, 1).text = "awa.diouf@exemple.sn"
    tableau.cell(1, 0).text = "Telephone"
    tableau.cell(1, 1).text = "+221 76 222 11 00"

    tampon = io.BytesIO()
    document.save(tampon)
    return tampon.getvalue()


@pytest.fixture(scope="session")
def csv_simple() -> bytes:
    contenu = (
        "nom,email,telephone,iban\n"
        "Mamadou FALL,mamadou.fall@exemple.sn,+221 77 555 44 33,"
        "SN91SN0100152000048500000765\n"
        "Awa Diouf,awa.diouf@exemple.sn,+221 76 222 11 00,\n"
    )
    return contenu.encode("utf-8")


@pytest.fixture(scope="session")
def csv_latin1() -> bytes:
    """Export de back-office en latin-1 : cas réel, doit être décodé."""
    return "nom,ville\nThérèse Kouamé,Abidjan\n".encode("latin-1")


@pytest.fixture(scope="session")
def image_jpeg() -> bytes:
    """Image de texte, rendue avec une police lisible par Tesseract."""
    import cv2
    import numpy as np

    image = np.full((520, 1100, 3), 255, dtype=np.uint8)
    lignes = [
        "ATTESTATION DE DOMICILE",
        "Nom: Mamadou FALL",
        "Ne le 03/07/1988 a Thies",
        "CNI 1988070312345",
        "Email mamadou.fall@exemple.sn",
        "Tel +221 77 555 44 33",
    ]
    for index, ligne in enumerate(lignes):
        cv2.putText(
            image,
            ligne,
            (40, 70 + index * 75),
            cv2.FONT_HERSHEY_SIMPLEX,
            1.1,
            (0, 0, 0),
            2,
            cv2.LINE_AA,
        )
    succes, encodee = cv2.imencode(".jpg", image, [int(cv2.IMWRITE_JPEG_QUALITY), 95])
    assert succes
    return encodee.tobytes()


@pytest.fixture(autouse=True)
def _sel_de_test(monkeypatch: pytest.MonkeyPatch) -> None:
    """Sel et backend NER déterministes pour tous les tests."""
    monkeypatch.setenv("HASH_SALT", "sel-de-test")
    monkeypatch.setenv("NER_BACKEND", "aucun")
    from app.config import get_settings

    get_settings.cache_clear()
    yield
    get_settings.cache_clear()
